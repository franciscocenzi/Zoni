import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def limpar_documento(doc):
    """Remove todo o corpo de texto (parágrafos e tabelas), mantendo Headers, Footers e Estilos Base."""
    for paragraph in reversed(doc.paragraphs):
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
        paragraph._p = paragraph._element = None
    
    for table in reversed(doc.tables):
        t_element = table._element
        t_element.getparent().remove(t_element)
        table._tbl = table._element = None

def aplicar_estilo_amigavel(doc):
    styles = doc.styles
    try:
        if 'ZoniHeading' not in styles:
            style = styles.add_style('ZoniHeading', WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles['Heading 1'] if 'Heading 1' in styles else styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)
            font.bold = True
            font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # azul corporativo
    except Exception:
        pass

def definir_largura_coluna(tabela, col_idx, largura_cm):
    """Define largura fixa para uma coluna específica."""
    from docx.shared import Cm
    for row in tabela.rows:
        if col_idx < len(row.cells):
            row.cells[col_idx].width = Cm(largura_cm)

def set_table_col_widths(table, widths_cm):
    """Define larguras fixas para todas as colunas da tabela (em cm)."""
    from docx.shared import Cm
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def adicionar_tabela_jinja(doc, cols, row_headers, list_var_name, fields_mapped,
                            style='Table Grid', widths_cm=None):
    """
    Cria uma tabela no Word com loop docxtpl para repetição de linhas.
    Usa {%tr for %} / {%tr endfor %} — sintaxe sem espaço após {% para docxtpl >= 0.9.
    """
    table = doc.add_table(rows=2, cols=cols)
    try:
        table.style = style
    except Exception:
        table.style = 'Table Grid'
    table.autofit = False

    # Cabeçalho
    hdr_cells = table.rows[0].cells
    for i, hc in enumerate(row_headers):
        hdr_cells[i].text = hc
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    def _render_field(field):
        """Envolve o campo em {{ }} a menos que já seja uma tag de bloco."""
        f = field.strip()
        if f.startswith('{%') or f.startswith('{{'):
            return field
        return '{{' + f + '}}'

    # {%tr for %} / {%tr endfor %} — repete a linha inteira (sintaxe docxtpl correta)
    row_cells = table.rows[1].cells
    row_cells[0].text = "{%tr for row in " + list_var_name + " %}" + _render_field(fields_mapped[0])
    for i in range(1, cols):
        if i == cols - 1:
            row_cells[i].text = _render_field(fields_mapped[i]) + "{%tr endfor %}"
        else:
            row_cells[i].text = _render_field(fields_mapped[i])

    # Fonte menor nas células de dados
    for cell in table.rows[1].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

    if widths_cm:
        set_table_col_widths(table, widths_cm)

    return table


def construir_modelo_zoni(doc):
    aplicar_estilo_amigavel(doc)
    style_h = 'ZoniHeading' if 'ZoniHeading' in doc.styles else 'Heading 1'
    
    # Título Principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("RELATÓRIO TÉCNICO DE ANÁLISE URBANÍSTICA")
    r.bold = True
    r.font.size = Pt(13)
    r.underline = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run("{{ TIPO_ANALISE }}")
    r2.bold = True
    r2.font.size = Pt(11)
    
    doc.add_paragraph("Emissão: {{ DATA_COMPLETA }} às {{ HORA }}", style='Normal')
    doc.add_paragraph()
    
    # 1. DADOS CADASTRAIS
    doc.add_paragraph("1. DADOS CADASTRAIS", style=style_h)
    adicionar_tabela_jinja(doc, 5, 
        ["Proprietário", "Inscrição / Matrícula", "Endereço", "Loteamento / Qd / Lt", "Área (m²)"],
        "DADOS_CADASTRAIS_LIST",
        ["row.proprietario", "row.inscricao", "row.endereco", "row.loteamento", "row.area"],
        widths_cm=[5.0, 3.5, 5.0, 4.0, 2.5]
    )
    doc.add_paragraph()
    
    # 2. LIMITES DO TERRENO
    doc.add_paragraph("2. LIMITES DO TERRENO (TESTADAS E DIVISAS)", style=style_h)
    adicionar_tabela_jinja(doc, 2, 
        ["Limite / Logradouro", "Comprimento (m)"],
        "TABELA_TESTADAS_LIST",
        ["row.limite", "row.comprimento"],
        widths_cm=[13.0, 4.0]
    )
    doc.add_paragraph()
    
    # 3. ZONEAMENTO E ÍNDICES
    doc.add_paragraph("3. ZONEAMENTO INCIDENTE", style=style_h)
    adicionar_tabela_jinja(doc, 10, 
        ["Zona", "Área (m²)", "%", "CA máx", "CA bas", "TPS", "TOS", "Pav Bas", "Pav Máx", "Recuo Fr"],
        "TABELA_ZONAS_LIST",
        ["row.codigo", "row.area", "row.perc", "row.ca_max", "row.ca_bas", "row.tps", "row.tos", "row.np_bas", "row.np_max", "row.rf"],
        widths_cm=[2.5, 2.2, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 2.0]
    )
    doc.add_paragraph()
    
    # 4. ÁREAS DE PRESERVAÇÃO PERMANENTE (APP)
    doc.add_paragraph("4. ÁREAS DE PRESERVAÇÃO PERMANENTE (APP)", style=style_h)
    doc.add_paragraph("Faixa NUIC: {{ APP_FAIXA_STATUS }} ({{ APP_FAIXA_LARGURA }} m) — {{ APP_FAIXA_OBS }}")
    doc.add_paragraph("Manguezal: {{ APP_MANGUE_STATUS }} — {{ APP_MANGUE_OBS }}")
    doc.add_paragraph()
    
    # 5. RISCOS GEOAMBIENTAIS — tabela sem bordas (Normal Table) com quadradinhos de cor
    doc.add_paragraph("5. RISCOS GEOAMBIENTAIS", style=style_h)
    t_r = doc.add_table(rows=3, cols=4)
    try:
        t_r.style = 'Normal Table'
    except Exception:
        t_r.style = 'Table Grid'
    t_r.autofit = False

    # Linha 0: cabeçalhos mesclados
    t_r.rows[0].cells[0].merge(t_r.rows[0].cells[1])
    t_r.rows[0].cells[0].text = "Suscetibilidade a Inundação"
    t_r.rows[0].cells[2].merge(t_r.rows[0].cells[3])
    t_r.rows[0].cells[2].text = "Suscetibilidade a Movimentos de Massa"
    for cell in [t_r.rows[0].cells[0], t_r.rows[0].cells[2]]:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # Linha 1: quadradinho colorido + texto do grau
    t_r.rows[1].cells[0].text = "{%cellbg RISCO_INUND_COR%}  "   # célula colorida
    t_r.rows[1].cells[1].text = "{{ RISCO_INUND_GRAU }}"
    t_r.rows[1].cells[2].text = "{%cellbg RISCO_MOV_COR%}  "    # célula colorida
    t_r.rows[1].cells[3].text = "{{ RISCO_MOV_GRAU }}"

    # Linha 2: recomendações mescladas
    t_r.rows[2].cells[0].merge(t_r.rows[2].cells[1])
    t_r.rows[2].cells[0].text = "{{ RISCO_INUND_RECOM }}"
    t_r.rows[2].cells[2].merge(t_r.rows[2].cells[3])
    t_r.rows[2].cells[2].text = "{{ RISCO_MOV_RECOM }}"

    set_table_col_widths(t_r, [1.0, 7.0, 1.0, 7.0])
    doc.add_paragraph()
    
    # 6. INCLINAÇÃO DO TERRENO — colunas de largura fixa
    doc.add_paragraph("6. INCLINAÇÃO DO TERRENO", style=style_h)
    adicionar_tabela_jinja(doc, 5, 
        ["Faixa de Inclinação", "Cor", "Área (m²)", "% da Área", "Notas"],
        "TABELA_INCLINACAO_LIST",
        ["row.faixa", "{%cellbg row.cor%}  ", "row.area", "row.perc", "row.notas"],
        style="Normal Table",
        widths_cm=[4.0, 1.2, 3.0, 2.5, 2.5]
    )
    doc.add_paragraph()
    
    # 7. NOTAS E CONDICIONANTES TÉCNICAS
    doc.add_paragraph("7. NOTAS E CONDICIONANTES TÉCNICAS", style=style_h)
    doc.add_paragraph("Notas técnicas e legislativas:").runs[0].font.bold = True
    doc.add_paragraph("{{ LISTA_NOTAS_ANEXO_BULLETS }}")
    doc.add_paragraph("Condicionantes para a análise:").runs[0].font.bold = True
    doc.add_paragraph("{{ LISTA_CONDICIONANTES_BULLETS }}")
    doc.add_paragraph("Restrições e pendências:").runs[0].font.bold = True
    doc.add_paragraph("{{ LISTA_RESTRICOES_BULLETS }}")
    doc.add_paragraph()

    # 8. MAPAS / IMAGENS
    doc.add_paragraph("8. MAPA DE SITUAÇÃO / ANEXOS GRÁFICOS", style=style_h)
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run("{{ IMAGEM_MAPA }}").italic = True

def main():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    out_dir = os.path.join(base_dir, "infraestrutura", "relatorios", "modelos")
    os.makedirs(out_dir, exist_ok=True)
    out_file = os.path.join(out_dir, "modelo_relatorio.docx")
    
    input_file = out_file if os.path.exists(out_file) else None
    print(f"Lendo base corporativa: {input_file}")
    doc = Document(input_file)
    limpar_documento(doc)
    construir_modelo_zoni(doc)
    doc.save(out_file)
    print(f"Modelo salvo em: {out_file}")

if __name__ == "__main__":
    main()
