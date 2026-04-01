"""
Renderizador DOCX para Zôni v2.
Estilo fiel ao modelo corporativo: só bordas horizontais, texto centralizado,
cabeçalhos em negrito, notas abaixo das tabelas com *.
"""
import os
import sys
import subprocess
import re
from datetime import datetime
from typing import List, Optional

from qgis.core import QgsMessageLog, Qgis

# ──────────────────────────────────────────────────────────────────
# Dependências
# ──────────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    QgsMessageLog.logMessage("python-docx ausente. Instalando...", "Zôni v2", Qgis.Warning)
    try:
        python_exe = os.path.join(sys.prefix, 'python.exe')
        if not os.path.exists(python_exe):
            python_exe = "python"
        subprocess.check_call([python_exe, "-m", "pip", "install", "python-docx"])
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except Exception as e:
        QgsMessageLog.logMessage(f"Falha ao instalar python-docx: {e}", "Zôni v2", Qgis.Critical)


# ──────────────────────────────────────────────────────────────────
# Estilo
# ──────────────────────────────────────────────────────────────────
FONT_NAME    = "Arial"
FONT_BODY    = Pt(9)
FONT_HDR_SEC = Pt(10)
FONT_TABLE   = Pt(8)
FONT_NOTA    = Pt(7)

COR_PRETO    = RGBColor(0x00, 0x00, 0x00)
COR_CINZA    = RGBColor(0x40, 0x40, 0x40)
COR_CINZA_CL = RGBColor(0x80, 0x80, 0x80)

# Largura útil A4 com margens corporativas (left≈3cm right≈2.5cm)
LARGURA_UTIL = 15.5   # cm


# ──────────────────────────────────────────────────────────────────
# Helpers de formatação de números
# ──────────────────────────────────────────────────────────────────
def _s(v) -> str:
    if v is None or str(v).strip() in ("", "None", "none"):
        return "-"
    return str(v).strip()

def _ff(v, dec=2) -> str:
    """Formata float PT-BR."""
    try:
        f = float(str(v).replace(",", "."))
        if f == int(f) and dec == 0:
            return f"{int(f):,}".replace(",", ".")
        return f"{f:,.{dec}f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return _s(v)

def _fp(v) -> str:
    """Formata percentual."""
    try:
        f = float(str(v).replace(",", ".").replace("%", ""))
        if 0 < f <= 1.0:
            f *= 100
        return f"{f:.1f}%".replace(".", ",")
    except Exception:
        return _s(v)


# ──────────────────────────────────────────────────────────────────
# Manipulação de XML de tabela
# ──────────────────────────────────────────────────────────────────
def _border_elem(val, sz="4", color="auto"):
    el = OxmlElement('w:border_placeholder')  # será substituído
    el.set(qn('w:val'), val)
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '0')
    el.set(qn('w:color'), color)
    return el

def set_table_borders(table, color="000000"):
    """
    Define apenas bordas HORIZONTAIS (topo, inferio e entre linhas).
    Verticais ficam invisíveis — estilo do modelo corporativo.
    """
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove borders existentes
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)

    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        if side in ('left', 'right', 'insideV'):
            el.set(qn('w:val'), 'none')
            el.set(qn('w:sz'), '0')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), 'auto')
        else:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)

def set_cell_color(cell, hex_color: str):
    """Cor de fundo de célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.replace("#", "").upper())
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    tcPr.append(shd)

def center_cell(cell):
    """Centraliza todos os parágrafos de uma célula."""
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ──────────────────────────────────────────────────────────────────
# Helpers de composição de documento
# ──────────────────────────────────────────────────────────────────
def _run(para, texto, bold=False, italic=False, size=None, color=None, font=FONT_NAME):
    r = para.add_run(texto)
    r.bold = bold
    r.italic = italic
    r.font.name = font
    r.font.size = size or FONT_BODY
    r.font.color.rgb = color or COR_PRETO
    return r

def _add_heading_sec(doc, texto: str):
    """Título de seção: negrito, cinza escuro, sem decoração especial."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    _run(p, texto, bold=True, size=FONT_HDR_SEC, color=COR_CINZA)
    return p

def _add_table_title(doc, titulo: str):
    """Título centralizado da tabela (como no modelo)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    _run(p, titulo, bold=True, size=Pt(9), color=COR_CINZA)

def _add_nota(doc, texto: str, simbolo="*"):
    """Nota abaixo da tabela em tamanho menor com símbolo de referência."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(4)
    _run(p, f"{simbolo} {texto}", size=FONT_NOTA, color=COR_CINZA_CL)


def _add_table(doc, titulo: str, headers: List[str], rows_data: List[List[str]],
               widths_cm: List[float], notas: List[str] = None,
               hdr_bold=True, zebra=False):
    """
    Tabela estilo modelo: só bordas horizontais, tudo centralizado,
    cabeçalho em negrito, larguras fixas.
    """
    if titulo:
        _add_table_title(doc, titulo)

    n_cols = len(headers)
    n_rows = max(2, 1 + len(rows_data))
    table  = doc.add_table(rows=n_rows, cols=n_cols)
    table.style  = 'Table Grid'
    table.autofit = False

    # Bordas só horizontais
    set_table_borders(table)

    # Cabeçalho
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        for para in cell.paragraphs:
            para.clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, h, bold=hdr_bold, size=FONT_TABLE)
        if zebra:
            set_cell_color(cell, "EFEFEF")

    # Dados
    for ri, row_vals in enumerate(rows_data):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_vals[:n_cols]):
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, _s(val), size=FONT_TABLE)
            if zebra and ri % 2 == 0:
                set_cell_color(cell, "FAFAFA")

    # Larguras fixas
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(widths_cm):
                cell.width = Cm(widths_cm[ci])

    # Notas abaixo
    if notas:
        for i, nota in enumerate(notas):
            simb = "*" if len(notas) == 1 else ("*" * (i + 1))
            _add_nota(doc, nota, simbolo=simb)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    return table


# ──────────────────────────────────────────────────────────────────
# Seções do relatório
# ──────────────────────────────────────────────────────────────────
def _sec_dados_cadastrais(doc, ctx):
    _add_heading_sec(doc, "I. DADOS CADASTRAIS")
    ident = ctx.get("identificacao") or {}
    ident_list = ident if isinstance(ident, list) else [ident]
    rows = []
    for d in ident_list:
        insc = _s(d.get("inscricao_imobiliaria"))
        cad  = _s(d.get("numero_cadastral"))
        ids  = " / ".join(x for x in [insc, cad] if x != "-") or "-"
        logr = _s(d.get("logradouro"))
        num  = _s(d.get("numero", "S/N"))
        bairro = _s(d.get("bairro"))
        end_parts = []
        if logr != "-": end_parts.append(f"{logr}, {num}")
        if bairro != "-": end_parts.append(bairro)
        endereco = " — ".join(end_parts) or "-"
        lot  = _s(d.get("loteamento"))
        qd   = _s(d.get("quadra"))
        lt   = _s(d.get("lote"))
        lot_str = " | ".join(x for x in [
            lot if lot != "-" else None,
            f"Qd {qd}" if qd != "-" else None,
            f"Lt {lt}" if lt != "-" else None
        ] if x) or "-"
        area = _ff(d.get("area_m2")) if d.get("area_m2") else "-"
        rows.append([_s(d.get("proprietario")), ids, endereco, lot_str, area])

    # 3.5+2.5+4.0+3.0+2.5 = 15.5
    _add_table(doc, "Dados Cadastrais do(s) Imóvel(is)",
               ["Proprietário", "Inscrição / Cad.", "Endereço", "Loteamento / Qd / Lt", "Área (m²)"],
               rows, widths_cm=[3.5, 2.5, 4.0, 3.0, 2.5])


def _sec_testadas(doc, ctx):
    _add_heading_sec(doc, "II. LIMITES DO TERRENO")
    segs = ctx.get("segmentos_limites") or []
    testadas_log = ctx.get("testadas_por_logradouro") or {}
    confrontantes = ctx.get("confrontantes_por_proprietario") or {}
    rows = []
    if isinstance(testadas_log, dict) and testadas_log:
        for log, comp in testadas_log.items():
            rows.append([f"Testada — {log}", _ff(comp)])
    if isinstance(confrontantes, dict) and confrontantes:
        for prop, comp in confrontantes.items():
            rows.append([f"Divisa — {prop}", _ff(comp)])
    if not rows and isinstance(segs, list):
        for s in segs:
            tipo = (s.get("tipo_limite") or "").upper()
            log  = _s(s.get("logradouro"))
            conf = _s(s.get("confrontante"))
            comp = _ff(s.get("comprimento_m"))
            desc = f"Testada — {log}" if tipo == "TESTADA" else f"Divisa — {conf}"
            rows.append([desc, comp])

    # 11.5+4.0 = 15.5
    _add_table(doc, "Testadas e Divisas",
               ["Limite / Logradouro / Confrontante", "Comprimento (m)"],
               rows, widths_cm=[11.5, 4.0])


def _sec_zoneamento(doc, ctx):
    _add_heading_sec(doc, "III. ZONEAMENTO INCIDENTE")
    zr    = ctx.get("zoneamento_resolvido") or {}
    zonas = zr.get("zonas") or []
    rows  = []
    notas_z = []

    if zonas:
        for z in zonas:
            param  = z.get("parametros") or {}
            extras = param.get("extras") or {}
            rows.append([
                _s(z.get("codigo")),
                _ff(z.get("area_m2")),
                _ff(z.get("percentual_area"), dec=1),
                _ff(param.get("CA_max")),
                _ff(param.get("CA_bas")),
                _fp(param.get("Tperm")),
                _fp(param.get("Tocup")),
                _s(param.get("Npav_bas")),
                _s(param.get("Npav_max")),
                _s(extras.get("RF") or extras.get("RF_Sec")),
            ])
        obs = zr.get("resumo") or ""
        for o in (zr.get("observacoes") or []):
            notas_z.append(str(o))
        if obs:
            notas_z.insert(0, obs)
    else:
        z_nome  = (ctx.get("zoneamento") or {}).get("zona", "-")
        indices = ctx.get("indices") or {}
        param   = indices.get("parametros") or {}
        extras  = param.get("extras") or {}
        rows.append([
            z_nome, _ff(ctx.get("area_lote_m2")), "100",
            _ff(param.get("CA_max")), _ff(param.get("CA_bas")),
            _fp(param.get("Tperm")), _fp(param.get("Tocup")),
            _s(param.get("Npav_bas")), _s(param.get("Npav_max")),
            _s(extras.get("RF") or extras.get("RF_Sec")),
        ])
        for m in ((ctx.get("zoneamento") or {}).get("mensagens") or []):
            notas_z.append(str(m))

    # 2.0+2.0+1.1+1.2+1.2+1.2+1.2+1.3+1.3+2.0 = 15.5
    _add_table(doc, "Parâmetros Urbanísticos por Zona",
               ["Zona", "Área (m²)", "%", "CA máx", "CA bas", "TPS", "TOS",
                "Pav Bas", "Pav Máx", "Recuo Fr"],
               rows,
               widths_cm=[2.0, 2.0, 1.1, 1.2, 1.2, 1.2, 1.2, 1.3, 1.3, 2.0],
               notas=notas_z or None)


def _sec_app(doc, ctx):
    _add_heading_sec(doc, "IV. ÁREAS DE PRESERVAÇÃO PERMANENTE (APP)")
    amb       = ctx.get("ambiente") or {}
    em_nuic   = amb.get("em_app_faixa_nuic")
    largura   = _s(amb.get("largura_faixa_m"))
    notas_amb = amb.get("notas") or []

    em_mangue = amb.get("em_app_manguezal")
    status_n  = "Presente" if em_nuic   else "Ausente"
    status_m  = "Presente" if em_mangue else "Ausente"
    larg_str  = f"{largura} m" if em_nuic and largura != "-" else "-"

    obs_nuic   = "; ".join(str(n) for n in notas_amb[:2]) if notas_amb else ""
    obs_mangue = "; ".join(str(n) for n in notas_amb[2:4]) if len(notas_amb) > 2 else ""

    notas_tbl = []
    if obs_nuic:   notas_tbl.append(f"Faixa NUIC: {obs_nuic}")
    if obs_mangue: notas_tbl.append(f"Manguezal: {obs_mangue}")

    # 4.0+3.0+8.5 = 15.5
    _add_table(doc, "Situação das APPs",
               ["Tipo de APP", "Situação / Largura", "Observações"],
               [
                   ["Faixa Marginal (NUIC)", f"{status_n} — {larg_str}",
                    "Lote intersecta faixa de curso d'água." if em_nuic else "Sem curso d'água identificado."],
                   ["Manguezal", status_m,
                    "Manguezal detectado na área." if em_mangue else "Sem manguezal identificado."],
               ],
               widths_cm=[4.0, 3.0, 8.5],
               notas=notas_tbl or None)


def _sec_risco(doc, ctx):
    _add_heading_sec(doc, "V. RISCOS GEOAMBIENTAIS")
    risco = ctx.get("risco") or {}
    cl_i  = _s(risco.get("classe_inundacao"))
    cl_m  = _s(risco.get("classe_movimento_massa"))

    COR_MAP = {
        "ALTA":       "FFCCCC", "ALTO":       "FFCCCC",
        "MÉDIA":      "FFF3CD", "MEDIA":      "FFF3CD",
        "BAIXA":      "D4EDDA", "BAIXO":      "D4EDDA",
        "MUITO BAIXA":"D4EDDA",
    }
    RECOM_MAP = {
        "ALTA":       "Exige Estudo Hidrológico/Geotécnico completo e medidas de contenção.",
        "ALTO":       "Exige Estudo Hidrológico/Geotécnico completo e medidas de contenção.",
        "MÉDIA":      "Recomenda-se investigação geotécnica/hidrológica preliminar.",
        "MEDIA":      "Recomenda-se investigação geotécnica/hidrológica preliminar.",
        "BAIXA":      "Procedimentos construtivos padrão geralmente suficientes.",
        "BAIXO":      "Procedimentos construtivos padrão geralmente suficientes.",
        "MUITO BAIXA":"Muito baixa suscetibilidade identificada.",
    }

    def _info(classe):
        s = str(classe).upper()
        for k in COR_MAP:
            if k in s:
                return COR_MAP[k], classe, RECOM_MAP[k]
        return "EEEEEE", "Não classificado" if classe == "-" else classe, "-"

    cor_i, grau_i, recom_i = _info(cl_i)
    cor_m, grau_m, recom_m = _info(cl_m)

    # Tabela 2 colunas: inundação | movimento de massa
    # Linha 0: categorias (cabeçalho)
    # Linha 1: cor + grau lado a lado
    # Linha 2: recomendação

    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    set_table_borders(table)

    # Linha 0 — cabeçalhos
    table.rows[0].cells[0].text = ""
    table.rows[0].cells[1].text = ""
    p0 = table.rows[0].cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p0, "Suscetibilidade a Inundação", bold=True, size=FONT_TABLE)
    p1 = table.rows[0].cells[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p1, "Suscetibilidade a Movimentos de Massa", bold=True, size=FONT_TABLE)

    # Linha 1 — célula de cor + grau
    # Usamos uma sub-tabela 1x2 dentro de cada célula principal
    # para ter o quadradinho colorido ao lado do texto — usando 2 colunas extras
    # Mais simples: tabela de 4 colunas, colunas 0 e 2 são as de cor
    table.rows[1].cells[0].text = ""
    c10 = table.rows[1].cells[0].paragraphs[0]
    c10.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(c10, f"  {grau_i}  ", bold=True, size=Pt(10))
    set_cell_color(table.rows[1].cells[0], cor_i)

    table.rows[1].cells[1].text = ""
    c11 = table.rows[1].cells[1].paragraphs[0]
    c11.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(c11, f"  {grau_m}  ", bold=True, size=Pt(10))
    set_cell_color(table.rows[1].cells[1], cor_m)

    # Linha 2 — recomendações
    table.rows[2].cells[0].text = ""
    c20 = table.rows[2].cells[0].paragraphs[0]
    c20.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(c20, recom_i, size=FONT_TABLE)

    table.rows[2].cells[1].text = ""
    c21 = table.rows[2].cells[1].paragraphs[0]
    c21.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(c21, recom_m, size=FONT_TABLE)

    # Larguras 7.75+7.75 = 15.5
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = Cm(7.75)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _sec_inclinacao(doc, ctx):
    _add_heading_sec(doc, "VI. INCLINAÇÃO DO TERRENO")
    incl   = ctx.get("inclinacao") or {}
    faixas = incl.get("faixas") or [] if isinstance(incl, dict) else []
    notas_inc = []

    if not faixas:
        msg = incl.get("mensagem", "Análise de inclinação não disponível.") if isinstance(incl, dict) else "Não disponível."
        p = doc.add_paragraph(msg)
        p.runs[0].font.size = FONT_BODY
        doc.add_paragraph()
        return

    table = doc.add_table(rows=1 + len(faixas), cols=5)
    table.style = 'Table Grid'
    table.autofit = False
    set_table_borders(table)

    _add_table_title(doc, "Distribuição de Inclinação do Terreno")

    # Rebuild table after title (add_table_title adds paragraph, but table already created above)
    # Actually let me add title before:
    # I'll add title first then remake the table inline below

    hdrs = ["Faixa de Inclinação", "Cor", "Área (m²)", "% da Área", "Notas"]
    for i, h in enumerate(hdrs):
        c = table.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, h, bold=True, size=FONT_TABLE)

    for ri, f in enumerate(faixas):
        row = table.rows[ri + 1]
        cor = str(f.get("cor", "#CCCCCC")).replace("#", "").upper()
        nota_faixa = "APP" if f.get("app") else "-"
        if f.get("app"):
            notas_inc.append(f"Limite >{f.get('faixa','?')}: área de APP por inclinação.")
        vals = [_s(f.get("faixa")), "", _ff(f.get("area_m2")),
                _ff(f.get("percentual"), dec=1) + "%", nota_faixa]
        for ci, val in enumerate(vals):
            row.cells[ci].text = ""
            p = row.cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, val, size=FONT_TABLE)
        set_cell_color(row.cells[1], cor)

    # 4.5+1.5+3.5+3.5+2.5 = 15.5
    widths = [4.5, 1.5, 3.5, 3.5, 2.5]
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(widths):
                cell.width = Cm(widths[ci])

    if notas_inc:
        for i, nota in enumerate(notas_inc):
            _add_nota(doc, nota)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _sec_notas(doc, ctx):
    """Notas e condicionantes como rodapé de seção (não tabela)."""
    _add_heading_sec(doc, "VII. NOTAS E CONDICIONANTES TÉCNICAS")

    try:
        from infraestrutura.relatorios.renderizador_html import _montar_listas_notas_separadas
        listas = _montar_listas_notas_separadas(ctx)
    except Exception:
        listas = {}

    def _extrair(chave):
        html = listas.get(chave, "")
        if html and isinstance(html, str):
            return [re.sub('<[^>]+>', '', i).strip()
                    for i in re.findall(r'<li[^>]*>(.*?)</li>', html, re.S)
                    if i.strip()]
        return []

    secoes = [
        ("Notas Técnicas e Legislativas", _extrair("LISTA_NOTAS_ANEXO")),
        ("Condicionantes", _extrair("LISTA_CONDICIONANTES")),
        ("Restrições e Pendências", _extrair("LISTA_RESTRICOES")),
    ]

    has_any = any(itens for _, itens in secoes)
    if not has_any:
        p = doc.add_paragraph("Nenhuma nota ou condicionante registrada para este imóvel.")
        p.runs[0].font.size = FONT_BODY
        doc.add_paragraph()
        return

    for subtitulo, itens in secoes:
        if not itens:
            continue
        p = doc.add_paragraph()
        _run(p, f"{subtitulo}:", bold=True, size=FONT_BODY)
        for idx, item in enumerate(itens, 1):
            b = doc.add_paragraph(style='List Bullet')
            b.paragraph_format.left_indent = Cm(0.5)
            _run(b, item, size=FONT_TABLE)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _sec_mapa(doc):
    _add_heading_sec(doc, "VIII. MAPA DE SITUAÇÃO")
    p = doc.add_paragraph("[Mapa será inserido futuramente]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.size = FONT_TABLE
    p.runs[0].font.color.rgb = COR_CINZA_CL


# ──────────────────────────────────────────────────────────────────
# Classe principal
# ──────────────────────────────────────────────────────────────────
class RenderizadorDOCX:
    def __init__(self):
        self.base_dir    = os.path.dirname(os.path.abspath(__file__))
        self.modelo_path = os.path.join(self.base_dir, "modelos", "modelo_relatorio.docx")

    def renderizar_e_salvar(self, contexto: dict, caminho_saida: str) -> tuple:
        try:
            if os.path.exists(self.modelo_path):
                doc = Document(self.modelo_path)
                for p in list(doc.paragraphs):
                    p._element.getparent().remove(p._element)
                for t in list(doc.tables):
                    t._element.getparent().remove(t._element)
            else:
                doc = Document()

            agora = datetime.now()
            tipo  = "Gleba Unificada" if contexto.get("area_gleba_unificada") else "Lote"

            # Título do relatório
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, "RELATÓRIO TÉCNICO DE ANÁLISE URBANÍSTICA",
                 bold=True, size=Pt(13), color=COR_CINZA)

            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p2, tipo, bold=True, size=Pt(11), color=COR_CINZA)

            pe = doc.add_paragraph()
            _run(pe, f"Emissão: {agora.strftime('%d/%m/%Y')} às {agora.strftime('%H:%M')}",
                 size=FONT_NOTA, color=COR_CINZA_CL)
            doc.add_paragraph()

            _sec_dados_cadastrais(doc, contexto)
            _sec_testadas(doc, contexto)
            _sec_zoneamento(doc, contexto)
            _sec_app(doc, contexto)
            _sec_risco(doc, contexto)
            _sec_inclinacao(doc, contexto)
            _sec_notas(doc, contexto)
            _sec_mapa(doc)

            doc.save(caminho_saida)
            QgsMessageLog.logMessage(f"Relatório gerado: {caminho_saida}", "Zôni v2", Qgis.Success)
            return True, ""
        except Exception as e:
            QgsMessageLog.logMessage(f"Erro ao gerar DOCX: {e}", "Zôni v2", Qgis.Critical)
            import traceback
            QgsMessageLog.logMessage(traceback.format_exc(), "Zôni v2", Qgis.Critical)
            return False, str(e)
